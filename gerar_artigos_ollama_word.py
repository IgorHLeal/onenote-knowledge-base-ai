import json
import os
import re
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ============================================================
# DIRETÓRIOS
# ============================================================

BASE_DIR = os.path.dirname(
    os.path.abspath(__file__)
)

OLLAMA_RESPOSTAS_DIR = os.path.join(
    BASE_DIR,
    "output",
    "processamento",
    "ollama",
    "respostas"
)

ARTIGOS_DIR = os.path.join(
    BASE_DIR,
    "output",
    "artigos"
)

ERROS_DIR = os.path.join(
    ARTIGOS_DIR,
    "erros"
)

PROCEDIMENTOS_DIR = os.path.join(
    ARTIGOS_DIR,
    "procedimentos"
)

REGISTRO_CODIGOS_FILE = os.path.join(
    ARTIGOS_DIR,
    "registro_codigos.json"
)

RELATORIO_FILE = os.path.join(
    ARTIGOS_DIR,
    "relatorio_geracao_ollama.json"
)


# ============================================================
# CONFIGURAÇÕES
# ============================================================

ERRO_INICIAL = 1
PRC_INICIAL = 12

IDS_IGNORAR = {
    "65",
    "73",
    "76",
    "83",
    "102",
}

TIPOS_VALIDOS = {
    "ERRO",
    "PROCEDIMENTO",
}


# ============================================================
# JSON
# ============================================================

def carregar_json(
    caminho,
    padrao=None
):

    if not os.path.exists(
        caminho
    ):

        if padrao is not None:
            return padrao

        raise FileNotFoundError(
            f"Arquivo não encontrado: {caminho}"
        )

    with open(
        caminho,
        "r",
        encoding="utf-8-sig"
    ) as arquivo:

        return json.load(
            arquivo
        )


def salvar_json(
    dados,
    caminho
):

    os.makedirs(
        os.path.dirname(
            caminho
        ),
        exist_ok=True
    )

    temporario = (
        caminho
        + ".tmp"
    )

    with open(
        temporario,
        "w",
        encoding="utf-8"
    ) as arquivo:

        json.dump(
            dados,
            arquivo,
            ensure_ascii=False,
            indent=2
        )

    os.replace(
        temporario,
        caminho
    )


# ============================================================
# TEXTO
# ============================================================

def texto_valido(
    valor
):

    if valor is None:
        return None

    if isinstance(
        valor,
        str
    ):

        valor = valor.strip()

        return valor or None

    valor = str(
        valor
    ).strip()

    return valor or None


def lista_textos(
    valor
):

    if valor is None:
        return []

    if isinstance(
        valor,
        str
    ):

        valor = valor.strip()

        return (
            [valor]
            if valor
            else []
        )

    if not isinstance(
        valor,
        list
    ):

        valor = [
            valor
        ]

    resultado = []

    for item in valor:

        if item is None:
            continue

        if isinstance(
            item,
            dict
        ):
            continue

        texto = str(
            item
        ).strip()

        if not texto:
            continue

        if texto not in resultado:

            resultado.append(
                texto
            )

    return resultado


def limpar_pontuacao_final(
    texto
):

    texto = texto_valido(
        texto
    )

    if not texto:
        return None

    return texto.strip()


def nome_arquivo_seguro(
    nome
):

    nome = re.sub(
        r'[<>:"/\\|?*]',
        "-",
        nome
    )

    nome = re.sub(
        r"\s+",
        " ",
        nome
    )

    nome = nome.strip(
        " ."
    )

    if len(
        nome
    ) > 150:

        nome = nome[
            :150
        ].rstrip()

    return nome


# ============================================================
# ESTILO WORD
# ============================================================

def configurar_documento(
    documento
):

    secao = documento.sections[
        0
    ]

    secao.top_margin = Cm(
        2
    )

    secao.bottom_margin = Cm(
        2
    )

    secao.left_margin = Cm(
        2.5
    )

    secao.right_margin = Cm(
        2.5
    )

    normal = documento.styles[
        "Normal"
    ]

    normal.font.name = (
        "Arial"
    )

    normal.font.size = Pt(
        11
    )

    normal._element.rPr.rFonts.set(
        qn(
            "w:eastAsia"
        ),
        "Arial"
    )

    for nome in (
        "Title",
        "Heading 1",
        "Heading 2",
    ):

        estilo = documento.styles[
            nome
        ]

        estilo.font.name = (
            "Arial"
        )

        estilo._element.rPr.rFonts.set(
            qn(
                "w:eastAsia"
            ),
            "Arial"
        )

    documento.styles[
        "Title"
    ].font.size = Pt(
        16
    )

    documento.styles[
        "Title"
    ].font.bold = True

    documento.styles[
        "Heading 1"
    ].font.size = Pt(
        11
    )

    documento.styles[
        "Heading 1"
    ].font.bold = True

    documento.styles[
        "Heading 2"
    ].font.size = Pt(
        11
    )

    documento.styles[
        "Heading 2"
    ].font.bold = True


def adicionar_borda_inferior(
    paragrafo
):

    p = paragrafo._p

    pPr = p.get_or_add_pPr()

    pBdr = OxmlElement(
        "w:pBdr"
    )

    bottom = OxmlElement(
        "w:bottom"
    )

    bottom.set(
        qn(
            "w:val"
        ),
        "single"
    )

    bottom.set(
        qn(
            "w:sz"
        ),
        "6"
    )

    bottom.set(
        qn(
            "w:space"
        ),
        "2"
    )

    bottom.set(
        qn(
            "w:color"
        ),
        "BFBFBF"
    )

    pBdr.append(
        bottom
    )

    pPr.append(
        pBdr
    )


def adicionar_titulo(
    documento,
    texto
):

    p = documento.add_paragraph()

    p.alignment = (
        WD_ALIGN_PARAGRAPH.LEFT
    )

    run = p.add_run(
        texto
    )

    run.bold = True

    run.font.size = Pt(
        15
    )

    adicionar_borda_inferior(
        p
    )


def adicionar_rotulo_valor(
    documento,
    rotulo,
    valor
):

    valores = lista_textos(
        valor
    )

    if not valores:
        return

    p = documento.add_paragraph()

    rotulo_run = p.add_run(
        f"{rotulo}: "
    )

    rotulo_run.bold = True

    p.add_run(
        ", ".join(
            valores
        )
    )


def adicionar_secao_texto(
    documento,
    titulo,
    texto
):

    texto = texto_valido(
        texto
    )

    if not texto:
        return

    documento.add_heading(
        titulo,
        level=1
    )

    documento.add_paragraph(
        texto
    )


def adicionar_lista(
    documento,
    titulo,
    itens
):

    itens = lista_textos(
        itens
    )

    if not itens:
        return

    documento.add_heading(
        titulo,
        level=1
    )

    for item in itens:

        documento.add_paragraph(
            item,
            style="List Bullet"
        )


# ============================================================
# ETAPAS
# ============================================================

def adicionar_etapas(
    documento,
    titulo_principal,
    etapas
):

    if not isinstance(
        etapas,
        list
    ):

        return

    etapas_validas = [
        etapa
        for etapa in etapas
        if isinstance(
            etapa,
            dict
        )
        and texto_valido(
            etapa.get(
                "etapa"
            )
        )
    ]

    if not etapas_validas:
        return

    documento.add_heading(
        titulo_principal,
        level=1
    )

    for etapa in etapas_validas:

        titulo_etapa = texto_valido(
            etapa.get(
                "etapa"
            )
        )

        itens = lista_textos(
            etapa.get(
                "itens",
                []
            )
        )

        if not titulo_etapa:
            continue

        documento.add_heading(
            titulo_etapa,
            level=2
        )

        for item in itens:

            documento.add_paragraph(
                item,
                style="List Bullet"
            )

        # Subgrupos representam subdivisões internas da etapa.
        # Eles devem aparecer como subtítulos, nunca como bullets.
        subgrupos = etapa.get(
            "subgrupos",
            []
        )

        if not isinstance(
            subgrupos,
            list
        ):
            subgrupos = []

        for subgrupo in subgrupos:

            if not isinstance(
                subgrupo,
                dict
            ):
                continue

            titulo_subgrupo = texto_valido(
                subgrupo.get(
                    "titulo"
                )
            )

            itens_subgrupo = lista_textos(
                subgrupo.get(
                    "itens",
                    []
                )
            )

            if not titulo_subgrupo:
                continue

            # O subgrupo fica visualmente abaixo da etapa, mas sem criar
            # uma nova seção principal do documento.
            p_subgrupo = documento.add_paragraph()
            p_subgrupo.paragraph_format.left_indent = Cm(0.5)
            p_subgrupo.paragraph_format.space_before = Pt(4)
            p_subgrupo.paragraph_format.space_after = Pt(2)

            run_subgrupo = p_subgrupo.add_run(
                titulo_subgrupo
            )
            run_subgrupo.bold = True

            for item_subgrupo in itens_subgrupo:

                p_item = documento.add_paragraph(
                    item_subgrupo,
                    style="List Bullet"
                )
                p_item.paragraph_format.left_indent = Cm(1.0)


# ============================================================
# REGISTRO DE CÓDIGOS
# ============================================================

def carregar_registro_codigos():

    return carregar_json(
        REGISTRO_CODIGOS_FILE,
        padrao={
            "erros": {},
            "procedimentos": {},
        }
    )


def extrair_numero_codigo(
    codigo,
    prefixo
):

    if not codigo:
        return None

    match = re.fullmatch(
        rf"{prefixo}-(\d+)",
        codigo.strip(),
        flags=re.IGNORECASE
    )

    if not match:
        return None

    return int(
        match.group(
            1
        )
    )


def proximo_codigo(
    registro_codigos,
    tipo,
    id_interno
):

    id_interno = str(
        id_interno
    )

    if tipo == "ERRO":

        mapa = registro_codigos.setdefault(
            "erros",
            {}
        )

        prefixo = "ERRO"

        inicial = (
            ERRO_INICIAL
        )

    else:

        mapa = registro_codigos.setdefault(
            "procedimentos",
            {}
        )

        prefixo = "PRC"

        inicial = (
            PRC_INICIAL
        )

    if id_interno in mapa:

        return mapa[
            id_interno
        ]

    numeros = []

    for codigo in mapa.values():

        numero = extrair_numero_codigo(
            codigo,
            prefixo
        )

        if numero is not None:

            numeros.append(
                numero
            )

    numero = (
        max(
            numeros,
            default=(
                inicial - 1
            )
        )
        + 1
    )

    codigo = (
        f"{prefixo}-"
        f"{numero:03d}"
    )

    mapa[
        id_interno
    ] = codigo

    return codigo


# ============================================================
# RESPOSTAS DO OLLAMA
# ============================================================

def listar_respostas():

    if not os.path.exists(
        OLLAMA_RESPOSTAS_DIR
    ):

        raise FileNotFoundError(
            (
                "Diretório de respostas "
                "do Ollama não encontrado: "
                f"{OLLAMA_RESPOSTAS_DIR}"
            )
        )

    arquivos = []

    for nome in os.listdir(
        OLLAMA_RESPOSTAS_DIR
    ):

        if not re.fullmatch(
            r"artigo_\d+\.json",
            nome,
            flags=re.IGNORECASE
        ):

            continue

        arquivos.append(
            os.path.join(
                OLLAMA_RESPOSTAS_DIR,
                nome
            )
        )

    arquivos.sort()

    return arquivos


def carregar_pacote_ollama(
    caminho
):

    pacote = carregar_json(
        caminho
    )

    if not isinstance(
        pacote,
        dict
    ):

        raise ValueError(
            (
                "Formato inválido no arquivo: "
                f"{caminho}"
            )
        )

    return pacote


# ============================================================
# VALIDAÇÃO PARA WORD
# ============================================================

def validar_pacote(
    pacote
):

    problemas = []

    id_interno = str(
        pacote.get(
            "id_interno",
            ""
        )
    ).strip()

    if not id_interno:

        problemas.append(
            "id_interno ausente."
        )

    status = (
        pacote.get(
            "status"
        )
        or ""
    ).strip().upper()

    if status != "APROVADO":

        problemas.append(
            (
                "Status diferente de APROVADO: "
                f"{status or 'VAZIO'}."
            )
        )

    resposta = pacote.get(
        "resposta_final"
    )

    if not isinstance(
        resposta,
        dict
    ):

        problemas.append(
            "resposta_final ausente ou inválida."
        )

        return problemas

    tipo = (
        resposta.get(
            "tipo_artigo"
        )
        or
        pacote.get(
            "tipo_artigo"
        )
        or ""
    ).strip().upper()

    if tipo not in TIPOS_VALIDOS:

        problemas.append(
            (
                "tipo_artigo inválido: "
                f"{tipo or 'VAZIO'}."
            )
        )

    titulo = texto_valido(
        resposta.get(
            "titulo"
        )
    )

    if not titulo:

        problemas.append(
            "Título ausente."
        )

    if tipo == "ERRO":

        etapas = resposta.get(
            "como_resolver",
            []
        )

    else:

        etapas = resposta.get(
            "procedimento",
            []
        )

    if not isinstance(
        etapas,
        list
    ):

        problemas.append(
            "Campo de etapas inválido."
        )

    if (
        tipo == "PROCEDIMENTO"
        and
        not etapas
    ):

        problemas.append(
            "PROCEDIMENTO sem etapas."
        )

    return problemas


# ============================================================
# GERAR ERRO
# ============================================================

def gerar_documento_erro(
    resposta,
    codigo,
    caminho
):

    documento = Document()

    configurar_documento(
        documento
    )

    titulo = texto_valido(
        resposta.get(
            "titulo"
        )
    ) or "Sem título"

    adicionar_titulo(
        documento,
        (
            f"{codigo} - "
            f"{titulo}"
        )
    )

    adicionar_rotulo_valor(
        documento,
        "Produto",
        resposta.get(
            "produto",
            []
        )
    )

    adicionar_secao_texto(
        documento,
        "Objetivo",
        resposta.get(
            "objetivo"
        )
    )

    adicionar_lista(
        documento,
        "Como o cliente normalmente relata",
        resposta.get(
            "como_cliente_relata",
            []
        )
    )

    adicionar_lista(
        documento,
        "Causa provável",
        resposta.get(
            "causa_provavel",
            []
        )
    )

    adicionar_lista(
        documento,
        "Mensagens de erro relacionadas",
        resposta.get(
            "mensagens_erro",
            []
        )
    )

    adicionar_etapas(
        documento,
        "Como resolver",
        resposta.get(
            "como_resolver",
            []
        )
    )

    adicionar_lista(
        documento,
        "Quando escalar",
        resposta.get(
            "quando_escalar",
            []
        )
    )

    adicionar_lista(
        documento,
        "Evidências obrigatórias (quando escalonado)",
        resposta.get(
            "evidencias_obrigatorias",
            []
        )
    )

    adicionar_lista(
        documento,
        "Observações",
        resposta.get(
            "observacoes",
            []
        )
    )

    documento.save(
        caminho
    )


# ============================================================
# GERAR PROCEDIMENTO
# ============================================================

def gerar_documento_procedimento(
    resposta,
    codigo,
    caminho
):

    documento = Document()

    configurar_documento(
        documento
    )

    titulo = texto_valido(
        resposta.get(
            "titulo"
        )
    ) or "Sem título"

    adicionar_titulo(
        documento,
        (
            f"{codigo} - "
            f"{titulo}"
        )
    )

    adicionar_rotulo_valor(
        documento,
        "Produto",
        resposta.get(
            "produto",
            []
        )
    )

    adicionar_secao_texto(
        documento,
        "Objetivo",
        resposta.get(
            "objetivo"
        )
    )

    adicionar_secao_texto(
        documento,
        "Quando utilizar",
        resposta.get(
            "quando_utilizar"
        )
    )

    adicionar_etapas(
        documento,
        "Procedimento",
        resposta.get(
            "procedimento",
            []
        )
    )

    adicionar_lista(
        documento,
        "Evidências obrigatórias",
        resposta.get(
            "evidencias_obrigatorias",
            []
        )
    )

    adicionar_lista(
        documento,
        "Observações",
        resposta.get(
            "observacoes",
            []
        )
    )

    documento.save(
        caminho
    )


# ============================================================
# GERAR UM
# ============================================================

def gerar_um(
    pacote,
    registro_codigos
):

    id_interno = str(
        pacote.get(
            "id_interno"
        )
    )

    resposta = pacote[
        "resposta_final"
    ]

    tipo = (
        resposta.get(
            "tipo_artigo"
        )
        or
        pacote.get(
            "tipo_artigo"
        )
        or ""
    ).strip().upper()

    codigo = proximo_codigo(
        registro_codigos,
        tipo,
        id_interno
    )

    titulo = (
        texto_valido(
            resposta.get(
                "titulo"
            )
        )
        or "Sem título"
    )

    nome = nome_arquivo_seguro(
        (
            f"{codigo} - "
            f"{titulo}"
        )
    )

    if tipo == "ERRO":

        os.makedirs(
            ERROS_DIR,
            exist_ok=True
        )

        caminho = os.path.join(
            ERROS_DIR,
            f"{nome}.docx"
        )

        gerar_documento_erro(
            resposta,
            codigo,
            caminho
        )

    else:

        os.makedirs(
            PROCEDIMENTOS_DIR,
            exist_ok=True
        )

        caminho = os.path.join(
            PROCEDIMENTOS_DIR,
            f"{nome}.docx"
        )

        gerar_documento_procedimento(
            resposta,
            codigo,
            caminho
        )

    return {
        "id_interno":
            id_interno,

        "codigo":
            codigo,

        "tipo":
            tipo,

        "titulo":
            titulo,

        "arquivo":
            caminho
    }


# ============================================================
# MAIN
# ============================================================



# ============================================================
# MAIN
# ============================================================

def main():

    import argparse

    parser = argparse.ArgumentParser()

    parser.add_argument(
        "id_interno",
        nargs="?"
    )

    args = parser.parse_args()

    print()
    print("=" * 70)
    print(
        "GERADOR DE ARTIGOS WORD - OLLAMA"
    )
    print("=" * 70)

    arquivos = listar_respostas()

    # ========================================================
    # FILTRO POR ID
    # ========================================================

    if args.id_interno:

        id_alvo = str(
            args.id_interno
        ).strip()

        arquivos_filtrados = []

        for caminho_json in arquivos:

            try:

                pacote = carregar_pacote_ollama(
                    caminho_json
                )

                id_interno = str(
                    pacote.get(
                        "id_interno",
                        ""
                    )
                ).strip()

                if id_interno == id_alvo:

                    arquivos_filtrados.append(
                        caminho_json
                    )

            except Exception:

                continue

        arquivos = arquivos_filtrados

        if not arquivos:

            print()
            print(
                f"Nenhuma resposta encontrada "
                f"para o ID {id_alvo}."
            )

            return

        print(
            f"Gerando somente o ID: "
            f"{id_alvo}"
        )

    else:

        print(
            f"Respostas encontradas: "
            f"{len(arquivos)}"
        )

    print(
        "IDs reservados para revisão manual: "
        + ", ".join(
            sorted(
                IDS_IGNORAR,
                key=int
            )
        )
    )

    registro_codigos = (
        carregar_registro_codigos()
    )

    gerados = []

    ignorados = []

    invalidos = []

    erros_execucao = []

    for indice, caminho_json in enumerate(
        arquivos,
        start=1
    ):

        try:

            pacote = carregar_pacote_ollama(
                caminho_json
            )

            id_interno = str(
                pacote.get(
                    "id_interno",
                    ""
                )
            ).strip()

            titulo = (
                pacote.get(
                    "titulo_original"
                )
                or
                os.path.basename(
                    caminho_json
                )
            )

            print()
            print(
                (
                    f"[{indice}/"
                    f"{len(arquivos)}] "
                    f"ID {id_interno or '?'}"
                )
            )

            print(
                titulo
            )

            if (
                id_interno
                in IDS_IGNORAR
            ):

                print(
                    "    IGNORADO - revisão manual."
                )

                ignorados.append(
                    {
                        "id_interno":
                            id_interno,

                        "motivo":
                            "Revisão manual pendente."
                    }
                )

                continue

            problemas = validar_pacote(
                pacote
            )

            if problemas:

                print(
                    "    IGNORADO - pacote não aprovado."
                )

                for problema in problemas:

                    print(
                        f"      - {problema}"
                    )

                invalidos.append(
                    {
                        "id_interno":
                            id_interno,

                        "arquivo":
                            caminho_json,

                        "problemas":
                            problemas
                    }
                )

                continue

            resultado = gerar_um(
                pacote,
                registro_codigos
            )

            salvar_json(
                registro_codigos,
                REGISTRO_CODIGOS_FILE
            )

            gerados.append(
                resultado
            )

            print(
                (
                    f"    GERADO: "
                    f"{resultado['codigo']}"
                )
            )

            print(
                (
                    f"    ARQUIVO: "
                    f"{resultado['arquivo']}"
                )
            )

        except Exception as erro:

            print(
                f"    ERRO: {erro}"
            )

            erros_execucao.append(
                {
                    "arquivo":
                        caminho_json,

                    "erro":
                        str(
                            erro
                        )
                }
            )

    relatorio = {
        "executado_em":
            datetime.now().isoformat(),

        "filtro_id":
            args.id_interno,

        "total_respostas_encontradas":
            len(
                arquivos
            ),

        "total_gerados":
            len(
                gerados
            ),

        "total_ignorados_revisao_manual":
            len(
                ignorados
            ),

        "total_invalidos":
            len(
                invalidos
            ),

        "total_erros_execucao":
            len(
                erros_execucao
            ),

        "ids_revisao_manual":
            sorted(
                IDS_IGNORAR,
                key=int
            ),

        "gerados":
            gerados,

        "ignorados":
            ignorados,

        "invalidos":
            invalidos,

        "erros_execucao":
            erros_execucao,
    }

    salvar_json(
        relatorio,
        RELATORIO_FILE
    )

    print()
    print()
    print("=" * 70)
    print(
        "RESULTADO DA GERAÇÃO"
    )
    print("=" * 70)

    print(
        f"Arquivos Word gerados: "
        f"{len(gerados)}"
    )

    print(
        f"Ignorados para revisão manual: "
        f"{len(ignorados)}"
    )

    print(
        f"Pacotes inválidos/não aprovados: "
        f"{len(invalidos)}"
    )

    print(
        f"Erros de execução: "
        f"{len(erros_execucao)}"
    )

    print("=" * 70)


if __name__ == "__main__":
    main()



# def main():

    # print()
    # print("=" * 70)
    # print(
        # "GERADOR DE ARTIGOS WORD - OLLAMA"
    # )
    # print("=" * 70)

    # arquivos = listar_respostas()

    # print(
        # f"Respostas encontradas: "
        # f"{len(arquivos)}"
    # )

    # print(
        # "IDs reservados para revisão manual: "
        # + ", ".join(
            # sorted(
                # IDS_IGNORAR,
                # key=int
            # )
        # )
    # )

    # registro_codigos = (
        # carregar_registro_codigos()
    # )

    # gerados = []

    # ignorados = []

    # invalidos = []

    # erros_execucao = []

    # for indice, caminho_json in enumerate(
        # arquivos,
        # start=1
    # ):

        # try:

            # pacote = carregar_pacote_ollama(
                # caminho_json
            # )

            # id_interno = str(
                # pacote.get(
                    # "id_interno",
                    # ""
                # )
            # ).strip()

            # titulo = (
                # pacote.get(
                    # "titulo_original"
                # )
                # or
                # os.path.basename(
                    # caminho_json
                # )
            # )

            # print()
            # print(
                # (
                    # f"[{indice}/"
                    # f"{len(arquivos)}] "
                    # f"ID {id_interno or '?'}"
                # )
            # )

            # print(
                # titulo
            # )

            # if (
                # id_interno
                # in IDS_IGNORAR
            # ):

                # print(
                    # "    IGNORADO - revisão manual."
                # )

                # ignorados.append(
                    # {
                        # "id_interno":
                            # id_interno,

                        # "motivo":
                            # "Revisão manual pendente."
                    # }
                # )

                # continue

            # problemas = validar_pacote(
                # pacote
            # )

            # if problemas:

                # print(
                    # "    IGNORADO - pacote não aprovado."
                # )

                # for problema in problemas:

                    # print(
                        # f"      - {problema}"
                    # )

                # invalidos.append(
                    # {
                        # "id_interno":
                            # id_interno,

                        # "arquivo":
                            # caminho_json,

                        # "problemas":
                            # problemas
                    # }
                # )

                # continue

            # resultado = gerar_um(
                # pacote,
                # registro_codigos
            # )

            # salvar_json(
                # registro_codigos,
                # REGISTRO_CODIGOS_FILE
            # )

            # gerados.append(
                # resultado
            # )

            # print(
                # (
                    # f"    GERADO: "
                    # f"{resultado['codigo']}"
                # )
            # )

        # except Exception as erro:

            # print(
                # f"    ERRO: {erro}"
            # )

            # erros_execucao.append(
                # {
                    # "arquivo":
                        # caminho_json,

                    # "erro":
                        # str(
                            # erro
                        # )
                # }
            # )

    # relatorio = {
        # "executado_em":
            # datetime.now().isoformat(),

        # "total_respostas_encontradas":
            # len(
                # arquivos
            # ),

        # "total_gerados":
            # len(
                # gerados
            # ),

        # "total_ignorados_revisao_manual":
            # len(
                # ignorados
            # ),

        # "total_invalidos":
            # len(
                # invalidos
            # ),

        # "total_erros_execucao":
            # len(
                # erros_execucao
            # ),

        # "ids_revisao_manual":
            # sorted(
                # IDS_IGNORAR,
                # key=int
            # ),

        # "gerados":
            # gerados,

        # "ignorados":
            # ignorados,

        # "invalidos":
            # invalidos,

        # "erros_execucao":
            # erros_execucao,
    # }

    # salvar_json(
        # relatorio,
        # RELATORIO_FILE
    # )

    # print()
    # print()
    # print("=" * 70)
    # print(
        # "RESULTADO DA GERAÇÃO"
    # )
    # print("=" * 70)

    # print(
        # f"Arquivos Word gerados: "
        # f"{len(gerados)}"
    # )

    # print(
        # f"Ignorados para revisão manual: "
        # f"{len(ignorados)}"
    # )

    # print(
        # f"Pacotes inválidos/não aprovados: "
        # f"{len(invalidos)}"
    # )

    # print(
        # f"Erros de execução: "
        # f"{len(erros_execucao)}"
    # )

    # print()

    # print(
        # "Diretório ERROS:"
    # )

    # print(
        # ERROS_DIR
    # )

    # print()

    # print(
        # "Diretório PROCEDIMENTOS:"
    # )

    # print(
        # PROCEDIMENTOS_DIR
    # )

    # print()

    # print(
        # "Relatório:"
    # )

    # print(
        # RELATORIO_FILE
    # )

    # print("=" * 70)


# if __name__ == "__main__":
    # main()